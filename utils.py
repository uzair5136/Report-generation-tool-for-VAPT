import io
import base64
import tempfile
import os
import logging
import json
from datetime import datetime
# ReportLab imports for non-editable PDFs
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, PageBreak
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas
from reportlab.graphics.shapes import Drawing, Rect
from reportlab.graphics.charts.piecharts import Pie
from reportlab.graphics.charts.barcharts import VerticalBarChart

# WeasyPrint is disabled for Windows compatibility
# Uncomment this block if you have installed GTK libraries and want to use WeasyPrint
# try:
#     # WeasyPrint for HTML to PDF (requires GTK libraries)
#     from weasyprint import HTML, CSS
#     WEASYPRINT_AVAILABLE = True
# except ImportError:
#     logging.warning("WeasyPrint not available. HTML-to-PDF conversion disabled.")
#     WEASYPRINT_AVAILABLE = False

# For now, we'll use only ReportLab which doesn't require GTK
WEASYPRINT_AVAILABLE = False
logging.info("Using ReportLab for PDF generation (no GTK dependencies required)")

# DocX for editable Word documents
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml

def generate_report_pdf(report, vulnerabilities, client, stats, images):
    """Generate a PDF report based on the provided data that exactly matches the provided VAPT template"""
    buffer = io.BytesIO()
    
    # Custom PDF generation with headers/footers and proper styling
    class VAPTDocTemplate(SimpleDocTemplate):
        """Custom document template with header and footer"""
        def __init__(self, filename, client_name="Client", **kwargs):
            SimpleDocTemplate.__init__(self, filename, **kwargs)
            self.pageinfo = kwargs.get('pageinfo', '')
            self.client_name = client_name
        
        def afterPage(self):
            """
            This method is called by SimpleDocTemplate without arguments,
            but it needs to create a callback that receives canvas and doc
            """
            def _afterPage(canvas, doc):
                # Save the canvas state
                canvas.saveState()
                
                # Header section
                # Dark green header bar exactly matching the template (adjusted height and position)
                canvas.setFillColor(colors.HexColor('#0B5345'))  # Darker green color from template
                canvas.rect(0, doc.height + doc.topMargin, 
                          doc.width + doc.leftMargin + doc.rightMargin, 
                          0.6*inch, fill=True, stroke=False)
                
                # Add logo placeholder in top left of header (if available)
                # Add text in header
                canvas.setFont('Helvetica-Bold', 10)
                canvas.setFillColor(colors.white)
                # Position text higher in the header bar
                canvas.drawString(doc.leftMargin, 
                               doc.height + doc.topMargin + 0.32*inch, 
                               f"VAPT Report - {self.client_name}")
                
                # Footer section
                # Medium green footer bar exactly matching template
                canvas.setFillColor(colors.HexColor('#138D75'))  # Medium green from template
                canvas.rect(0, doc.bottomMargin - 0.5*inch, 
                          doc.width + doc.leftMargin + doc.rightMargin, 
                          0.45*inch, fill=True, stroke=False)
                
                # Add page number in footer - position exactly as in template
                canvas.setFont('Helvetica', 9)
                canvas.setFillColor(colors.white)
                text = f"Page {doc.page} of " + "{total}" # Placeholder for total pages
                canvas.drawRightString(doc.width + doc.rightMargin - 10, 
                                    doc.bottomMargin - 0.27*inch, text)
                
                # Copyright text in footer
                canvas.setFont('Helvetica', 8)
                canvas.setFillColor(colors.white)
                canvas.drawString(doc.leftMargin, 
                               doc.bottomMargin - 0.27*inch, 
                               "© 2025 CyHEX Infotech - Confidential Report")
                
                canvas.restoreState()
                
            return _afterPage
    
    # Use our custom template with green header/footer
    doc = VAPTDocTemplate(buffer, client_name=client.name, pagesize=letter, 
                      topMargin=72, bottomMargin=72, 
                      leftMargin=50, rightMargin=50)
    elements = []
    
    # Styles based on exact template specifications
    styles = getSampleStyleSheet()
    
    # Define custom styles precisely matching the template document colors and spacing
    # Italic style for captions and image descriptions
    styles.add(ParagraphStyle(
        name='ImageCaption',
        fontName='Helvetica-Oblique',
        fontSize=9,
        textColor=colors.gray,  # Match exact template color 
        spaceAfter=8,
        leading=12
    ))
    
    # Title style exactly matching the template (centered, bold, exact font size)
    title_style = ParagraphStyle(
        name='TitleStyle',
        fontName='Helvetica-Bold',
        fontSize=24,  # Adjusted to match template exactly
        alignment=1,  # Center alignment
        spaceAfter=0,
        spaceBefore=6,
        leading=30,  # Line spacing
        textColor=colors.black  # Exact color from template
    )
    
    subtitle_style = ParagraphStyle(
        name='SubtitleStyle',
        fontName='Helvetica-Bold', 
        fontSize=16,  # Adjusted to match template exactly
        alignment=1,  # Center alignment
        spaceAfter=6,
        spaceBefore=4,
        leading=20,  # Line spacing
        textColor=colors.black  # Exact color from template
    )
    
    # Section headings (left-aligned, smaller than title)
    section_title = ParagraphStyle(
        name='SectionTitle',
        fontName='Helvetica-Bold',
        fontSize=14,
        alignment=0,  # Left alignment
        spaceAfter=10,
        spaceBefore=14,
        leading=16,  # Line spacing
        textColor=colors.black  # Exact color from template
    )
    
    subsection_title = ParagraphStyle(
        name='SubsectionTitle',
        fontName='Helvetica-Bold',
        fontSize=12,
        alignment=0,  # Left alignment
        spaceAfter=8,
        spaceBefore=12,
        leading=14,  # Line spacing
        textColor=colors.black  # Exact color from template
    )
    
    # Body text style exactly matching template with improved spacing
    body_text = ParagraphStyle(
        name='BodyText',
        fontName='Helvetica',
        fontSize=11,
        alignment=4,  # Justified text for professional appearance
        spaceAfter=18,  # Further increased spacing after paragraphs
        spaceBefore=6,  # Further increased spacing before paragraphs
        leading=20,  # Further increased line spacing for better readability
        textColor=colors.black  # Exact color from template
    )
    
    # Bullet point style with reduced spacing
    bullet_point_style = ParagraphStyle(
        name='BulletPoint',
        fontName='Helvetica',
        fontSize=11,
        alignment=0,  # Left aligned for bullet points
        spaceAfter=10,  # Reduced spacing after paragraphs for compact bullet lists
        spaceBefore=2,  # Reduced spacing before paragraphs
        leading=14,  # Reduced line spacing for more compact appearance
        textColor=colors.black,  # Exact color from template
        leftIndent=20  # Indentation for bullet points
    )
    
    # Title Page - Precisely matching the provided template
    # Add company logo if available (placeholder for now)
    
    # Title with exact spacing and alignment
    elements.append(Spacer(1, 30))  # Top margin
    elements.append(Paragraph("Vulnerability Assessment", title_style))
    elements.append(Paragraph("and Penetration Testing", title_style))
    elements.append(Spacer(1, 100))  # Large spacing as in template
    
    elements.append(Paragraph(f"{report.domain_type} Application Security Report", subtitle_style))
    elements.append(Spacer(1, 10))
    elements.append(Paragraph("L1 Report", subtitle_style))
    elements.append(Spacer(1, 10))
    elements.append(Paragraph(f"{datetime.now().strftime('%d %B %Y')}", subtitle_style))
    elements.append(PageBreak())
    
    # Client Details on a separate page immediately after the title page
    elements.append(Paragraph("Client Details", styles['Heading1']))
    elements.append(Spacer(1, 10))
    
    client_data = [
        ["Organization", client.name],
        ["Contact Person", client.contact_email.split("@")[0] if client.contact_email else ""],
        ["Email id", client.contact_email if client.contact_email else ""]
    ]
    
    client_table = Table(client_data, colWidths=[150, 300])
    client_table.setStyle(TableStyle([
        # Header column styling (first column) to match template exactly
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#0B5345')),  # Dark green header matching template exactly
        ('TEXTCOLOR', (0, 0), (0, -1), colors.white),  # White text
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (0, -1), 11),  # Font size from template
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        # Value column styling to match template
        ('BACKGROUND', (1, 0), (1, -1), colors.HexColor('#E8F5E9')),  # Light green background matching template
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#138D75')),  # Green grid color to match template exactly
        ('FONTSIZE', (1, 0), (1, -1), 10),  # Font size from template
    ]))
    elements.append(client_table)
    elements.append(Spacer(1, 30))
    
    # Document Revision
    elements.append(Paragraph("Document Revision", subsection_title))
    elements.append(Spacer(1, 10))
    
    revision_data = [
        ["Version", f"{report.version}.0"],
        ["Date", datetime.now().strftime('%d %B %Y')],
        ["Submitted By", client.submitted_by or "Security Assessment Team"]
    ]
    
    revision_table = Table(revision_data, colWidths=[150, 300])
    revision_table.setStyle(TableStyle([
        # Header column styling (first column) to match template exactly
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#0B5345')),  # Dark green header matching template exactly
        ('TEXTCOLOR', (0, 0), (0, -1), colors.white),  # White text
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (0, -1), 11),  # Font size from template
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        # Value column styling to match template
        ('BACKGROUND', (1, 0), (1, -1), colors.HexColor('#E8F5E9')),  # Light green background matching template
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#138D75')),  # Green grid color to match template exactly
        ('FONTSIZE', (1, 0), (1, -1), 10),  # Font size from template
    ]))
    elements.append(revision_table)
    elements.append(Spacer(1, 30))
    
    # Test Performed Details
    elements.append(Paragraph("Test Performed Details", subsection_title))
    elements.append(Spacer(1, 10))
    
    test_data = [
        ["Testing done By", client.testing_done_by or "Security Assessment Team"],
        ["Reviewed By", client.reviewed_by or "Senior Security Analyst"],
        ["Date", f"{report.start_date.strftime('%d %B %Y')} - {report.end_date.strftime('%d %B %Y')}"],
        ["Version", f"{report.version}.0"]
    ]
    
    test_table = Table(test_data, colWidths=[150, 300])
    test_table.setStyle(TableStyle([
        # Header column styling (first column) to match template exactly
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#0B5345')),  # Dark green header matching template exactly
        ('TEXTCOLOR', (0, 0), (0, -1), colors.white),  # White text
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (0, -1), 11),  # Font size from template
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ('TOPPADDING', (0, 0), (-1, -1), 8),
        # Value column styling to match template
        ('BACKGROUND', (1, 0), (1, -1), colors.HexColor('#E8F5E9')),  # Light green background matching template
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#138D75')),  # Green grid color to match template exactly
        ('FONTSIZE', (1, 0), (1, -1), 10),  # Font size from template
    ]))
    elements.append(test_table)
    
    # Add page break to move to next page
    elements.append(PageBreak())
    
    # Table of Contents
    elements.append(Paragraph("Table of Contents", section_title))
    toc_data = [
        ["Table of Contents", "4"],
        ["1. Executive Summary", "5"],
        ["    1.1 Summary", "5"],
        ["    1.2 Approach", "5"],
        ["    1.3 Disclaimer", "5"],
        ["    1.4 Limitations", "6"],
        ["    1.5 OWASP TOP 10", "6"],
        ["    1.6 Vulnerability Scoring", "6"],
        ["2. Checklist", "7"],
        ["3. Scope", "10"],
        ["    3.1 Key Findings", "10"],
        ["    3.2 Vulnerability Graph", "11"],
        ["4. Findings", "12"],
        ["5. Conclusions", "13"],
        ["6. Tools Used", "13"]
    ]
    
    toc = Table(toc_data, colWidths=[400, 50])
    toc.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 11),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('ALIGN', (1, 0), (1, -1), 'RIGHT'),
    ]))
    elements.append(toc)
    elements.append(PageBreak())
    
    # Executive Summary
    elements.append(Paragraph("1. Executive Summary", title_style))
    elements.append(Paragraph("1.1 Summary", section_title))
    
    summary_text = f"""
    Our security team conducted a penetration test on {client.name}'s {report.domain_type}
    environment, starting on {report.start_date.strftime('%d %B %Y')}. This assessment, combining automated tools and
    manual checks, aimed to uncover technical weaknesses in the application that
    could be exploited. The report details the identified vulnerabilities, their severity,
    and provides recommendations to mitigate any security risks they might pose.
    
    A total of {stats['total']} vulnerabilities were identified, categorized as follows:
    - Critical: {stats['critical']}
    - High: {stats['high']}
    - Medium: {stats['medium']}
    - Low: {stats['low']}
    """
    elements.append(Paragraph(summary_text, body_text))
    
    # Approach
    elements.append(Paragraph("1.2 Approach", section_title))
    
    # Initial text for the approach section
    intro_text = "The approach for this assessment included the following activities:"
    elements.append(Paragraph(intro_text, body_text))
    
    # Define approach as bullet points with each on a new line
    approach_bullet_points = [
        "• Exploring various application functionalities to enumerate threat & vulnerability in alignment with Open Web Application Security Project (OWASP) Top 10 vulnerabilities.",
        "• Performing information gathering/fingerprinting to identify software used/its version, web server details, ports, and services open, etc.",
        "• Performing vulnerability scanning to identify common vulnerabilities in the application layer and by using Burp and various testing tools in the Kali Linux distribution in conjunction with a range of manual analysis.",
        "• It should be noted that customized payloads and attack vectors were configured in Burp Suite to further enhance the identification of weakness in the application.",
        "• Analyzing the automated scan results for any vulnerabilities and ease of exploitability and providing proof of concept where safe exploits are possible.",
        "• Post-Exploitation process will be performed once we get access to the device using identified vulnerabilities/exploits.",
        "• Reporting identified vulnerabilities and recommended solutions to mitigate them; for ease of mitigation activities for application support personnel/developers' further details of CWEs were added."
    ]
    
    # Add each bullet point as a separate paragraph with reduced spacing
    for point in approach_bullet_points:
        elements.append(Paragraph(point, bullet_point_style))
    
    elements.append(Spacer(1, 10))
    
    # Disclaimer
    elements.append(Paragraph("1.3 Disclaimer", section_title))
    disclaimer_text = """
    This report contains confidential information about the security of the client's systems. 
    Distribution of this report should be limited to authorized personnel only.
    The testing was performed during a limited timeframe and may not have identified all possible vulnerabilities.
    """
    elements.append(Paragraph(disclaimer_text, body_text))
    
    # Limitations
    elements.append(Paragraph("1.4 Limitations", section_title))
    elements.append(Paragraph(report.limitations or "Not specified", styles['Normal']))
    
    # OWASP TOP 10
    elements.append(Paragraph("1.5 OWASP TOP 10", section_title))
    
    # Introduction for OWASP TOP 10
    intro_text = "The OWASP Top 10 is a regularly updated report outlining security concerns for web application security. The top 10 vulnerabilities for 2021 are:"
    elements.append(Paragraph(intro_text, body_text))
    
    # Define OWASP TOP 10 as bullet points with each on a new line
    owasp_bullet_points = [
        "• A01:2021 – Broken Access Control",
        "• A02:2021 – Cryptographic Failures",
        "• A03:2021 – Injection",
        "• A04:2021 – Insecure Design",
        "• A05:2021 – Security Misconfiguration",
        "• A06:2021 – Vulnerable and Outdated Components", 
        "• A07:2021 – Identification and Authentication Failures",
        "• A08:2021 – Software and Data Integrity Failures",
        "• A09:2021 – Security Logging and Monitoring Failures",
        "• A10:2021 – Server-Side Request Forgery"
    ]
    
    # Add each bullet point as a separate paragraph with reduced spacing
    for point in owasp_bullet_points:
        elements.append(Paragraph(point, bullet_point_style))
    
    # Vulnerability Scoring
    elements.append(Paragraph("1.6 Vulnerability Scoring", section_title))
    scoring_text = "The Risk level is divided in four categories:"
    elements.append(Paragraph(scoring_text, body_text))
    elements.append(Spacer(1, 10))
    
    # Create vulnerability scoring table exactly matching the template with text wrapping
    # Each paragraph will be properly formatted inside the table cell
    severity_paragraphs = {
        "Critical": Paragraph("Critical vulnerabilities provide attackers with remote root or administrator capabilities. Malicious users have the ability to compromise the entire host. Easy to detect and exploit and result in large asset damage.", body_text),
        "High": Paragraph("Exploitation of the vulnerability discovered on the system can directly lead an attacker to information allowing them to gain privileged access (e.g., administrator or root) to the system. These issues are often difficult to detect and exploit but can result in large asset damage.", body_text),
        "Medium": Paragraph("The vulnerability discovered on the system can directly lead to an attacker gaining non-privileged access (e.g., as a standard user) to the system or the vulnerability provides access that can be leveraged within one step to gain administrator-level access. These issues are easy to detect and exploit, but typically result in small asset damage.", body_text),
        "Low": Paragraph("The vulnerability discovered on the system provides low-level, but sufficient data to the attacker that may be used to launch a more informed attack against the target environment. In addition, the vulnerability may indirectly lead to an attacker gaining some form of access to the system. These issues can be difficult to detect and exploit and typically result in small asset damage.", body_text)
    }
    
    scoring_data = [
        ["Severity", "DESCRIPTION"],
        ["Critical", severity_paragraphs["Critical"]],
        ["High", severity_paragraphs["High"]],
        ["Medium", severity_paragraphs["Medium"]],
        ["Low", severity_paragraphs["Low"]]
    ]
    
    # Calculate column widths for the table - adjust to match template exactly
    col_widths = [80, 440]
    scoring_table = Table(scoring_data, colWidths=col_widths)
    
    # Style the table to match the template exactly
    scoring_table.setStyle(TableStyle([
        # Header row styling
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0B5345')),  # Dark green header matching template exactly
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),  # White text in header
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),  # Exact font size from template
        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
        ('TOPPADDING', (0, 0), (-1, 0), 10),
        
        # Data rows styling
        ('BACKGROUND', (0, 1), (0, -1), colors.HexColor('#E8F5E9')),  # Light green background for first column matching template
        ('BACKGROUND', (1, 1), (1, -1), colors.HexColor('#E8F5E9')),  # Light green background for second column matching template
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#138D75')),  # Green grid color to match template exactly
        ('FONTSIZE', (0, 1), (-1, -1), 10),  # Exact font size from template
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),  # Vertically centered text
        
        # Set severity cell background colors
        ('BACKGROUND', (0, 1), (0, 1), colors.HexColor('#FF0000')),  # Red for Critical
        ('BACKGROUND', (0, 2), (0, 2), colors.HexColor('#FFA500')),  # Orange for High
        ('BACKGROUND', (0, 3), (0, 3), colors.HexColor('#FFFF00')),  # Yellow for Medium
        ('BACKGROUND', (0, 4), (0, 4), colors.HexColor('#7FFF00')),  # Green for Low
        
        # Set text color for severity cells to ensure visibility
        ('TEXTCOLOR', (0, 1), (0, 1), colors.white),  # White text for Critical row
        ('TEXTCOLOR', (0, 2), (0, 2), colors.black),  # Black text for High row
        ('TEXTCOLOR', (0, 3), (0, 3), colors.black),  # Black text for Medium row
        ('TEXTCOLOR', (0, 4), (0, 4), colors.black),  # Black text for Low row
    ]))
    
    elements.append(scoring_table)
    
    elements.append(PageBreak())
    
    # Checklist
    elements.append(Paragraph("2. Checklist", styles['Heading1']))
    elements.append(Paragraph("The following security testing checklist was used during the assessment:", styles['Normal']))
    elements.append(Spacer(1, 0.2*inch))
    
    # Create security testing checklist table
    if report.checklist_data:
        try:
            checklist_data = json.loads(report.checklist_data)
            
            # Create table header
            checklist_table_data = [
                ["Category", "Outcome"]
            ]
            
            # Add categories and outcomes to table
            for category in checklist_data["categories"]:
                checklist_table_data.append([
                    category["name"],
                    category["outcome"]
                ])
                
            # Create table with appropriate styling
            checklist_table = Table(checklist_table_data, colWidths=[4*inch, 2*inch])
            
            # Style the table similar to others in the document
            table_style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0B5345')),  # Dark green header
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#138D75')),  # Medium green grid lines
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#E8F5E9')),  # Light green background
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                ('ALIGN', (1, 0), (1, -1), 'CENTER'),
            ])
            
            # Add alternating row colors for better readability
            for i in range(1, len(checklist_table_data), 2):
                table_style.add('BACKGROUND', (0, i), (-1, i), colors.HexColor('#D5F5E3'))
            
            checklist_table.setStyle(table_style)
            elements.append(checklist_table)
        except Exception as e:
            logging.error(f"Error processing checklist data: {str(e)}")
            elements.append(Paragraph("Error loading checklist data.", styles['Normal']))
    else:
        elements.append(Paragraph("No checklist data available.", styles['Normal']))
    
    elements.append(PageBreak())
    
    # Scope
    elements.append(Paragraph("3. Scope", styles['Heading1']))
    elements.append(Paragraph(report.scope or "Not specified", styles['Normal']))
    
    # Key Findings
    elements.append(Paragraph("3.1 Key Findings", styles['Heading2']))
    findings_data = [
        ["Severity", "Count"],
        ["Critical", str(stats['critical'])],
        ["High", str(stats['high'])],
        ["Medium", str(stats['medium'])],
        ["Low", str(stats['low'])],
        ["Total", str(stats['total'])]
    ]
    
    findings_table = Table(findings_data, colWidths=[200, 100])
    findings_table.setStyle(TableStyle([
        # Header row styling to match template exactly
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0B5345')),  # Dark green header matching template exactly
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),  # White text in header
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 11),  # Exact font size from template
        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
        ('TOPPADDING', (0, 0), (-1, 0), 10),
        # Data rows styling to match template
        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#E8F5E9')),  # Light green background matching template
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#138D75')),  # Green grid color to match template exactly
        ('FONTSIZE', (0, 1), (-1, -1), 10),  # Exact font size from template
        ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
        ('TOPPADDING', (0, 1), (-1, -1), 8),
    ]))
    elements.append(findings_table)
    
    # Vulnerability Graph
    elements.append(Paragraph("3.2 Vulnerability Graph", styles['Heading2']))
    
    # Bar chart for vulnerabilities
    if stats['total'] > 0:
        drawing = Drawing(400, 200)
        bc = VerticalBarChart()
        bc.x = 50
        bc.y = 50
        bc.height = 125
        bc.width = 300
        bc.data = [[stats['critical'], stats['high'], stats['medium'], stats['low']]]
        bc.bars[0].fillColor = colors.HexColor('#138D75')  # Green bar color to match the template exactly
        bc.categoryAxis.categoryNames = ['Critical', 'High', 'Medium', 'Low']
        bc.valueAxis.valueMin = 0
        bc.valueAxis.valueMax = max(stats['critical'], stats['high'], stats['medium'], stats['low']) + 1
        bc.valueAxis.valueStep = 1
        drawing.add(bc)
        elements.append(drawing)
    
    elements.append(PageBreak())
    
    # Findings
    elements.append(Paragraph("4. Findings", styles['Heading1']))
    
    # Organize vulnerabilities into groups
    grouped_vulns = {}
    ungrouped_vulns = []
    
    for vuln in vulnerabilities:
        if vuln.group_id:
            group_id = vuln.group_id
            if group_id not in grouped_vulns:
                grouped_vulns[group_id] = {
                    "name": vuln.group.name,
                    "vulns": []
                }
            grouped_vulns[group_id]["vulns"].append(vuln)
        else:
            ungrouped_vulns.append(vuln)
    
    section_index = 1
    
    # Display grouped vulnerabilities
    for group_id, group_data in grouped_vulns.items():
        elements.append(Paragraph(f"4.{section_index} {group_data['name']}", styles['Heading2']))
        
        # Add introduction text for the vulnerability group
        group_intro = f"""This group contains {len(group_data['vulns'])} related vulnerabilities that affect similar components 
        or share common attack vectors. Addressing these vulnerabilities as a group can provide a more comprehensive 
        security enhancement, as they often have interconnected root causes."""
        elements.append(Paragraph(group_intro, body_text))
        
        # Sort vulnerabilities within group by severity
        severity_order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}
        sorted_vulns = sorted(group_data['vulns'], key=lambda x: severity_order.get(x.severity, 4))
        
        for i, vuln in enumerate(sorted_vulns, 1):
            # Add formatted vulnerability title with severity badge
            elements.append(Paragraph(f"4.{section_index}.{i} {vuln.title}", styles['Heading3']))
            
            # Create a table for vulnerability summary
            summary_data = [
                ["Severity", vuln.severity],
                ["CWE ID", vuln.cwe_id or "Not Specified"],
                ["Status", "Open"],
                ["Date Identified", vuln.created_at.strftime('%d %B %Y') if hasattr(vuln, 'created_at') and vuln.created_at else "Not recorded"]
            ]
            
            summary_table = Table(summary_data, colWidths=[100, 350])
            summary_table.setStyle(TableStyle([
                # Left column styling (category headers) to match template exactly
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#0B5345')),  # Dark green header matching template exactly
                ('TEXTCOLOR', (0, 0), (0, -1), colors.white),  # White text
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (0, -1), 10),  # Exact font size from template
                ('BOTTOMPADDING', (0, 0), (0, -1), 8),
                ('TOPPADDING', (0, 0), (0, -1), 8),
                # Value column styling to match template
                ('BACKGROUND', (1, 0), (1, -1), colors.HexColor('#E8F5E9')),  # Light green background matching template
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#138D75')),  # Green grid color to match template exactly
                ('FONTSIZE', (1, 0), (1, -1), 10),  # Exact font size from template
                # Set red background for "Critical" severity
                ('BACKGROUND', (1, 0), (1, 0), colors.HexColor('#FF0000') if vuln.severity == 'Critical' else colors.HexColor('#E8F5E9')),
                # Set orange background for "High" severity
                ('BACKGROUND', (1, 0), (1, 0), colors.HexColor('#FFA500') if vuln.severity == 'High' else colors.HexColor('#E8F5E9')),
                # Set yellow background for "Medium" severity
                ('BACKGROUND', (1, 0), (1, 0), colors.HexColor('#FFFF00') if vuln.severity == 'Medium' else colors.HexColor('#E8F5E9')),
                # Set green background for "Low" severity
                ('BACKGROUND', (1, 0), (1, 0), colors.HexColor('#7FFF00') if vuln.severity == 'Low' else colors.HexColor('#E8F5E9')),
            ]))
            elements.append(summary_table)
            elements.append(Spacer(1, 10))
            
            # Vulnerability details with improved formatting
            elements.append(Paragraph("<b>Description</b>", section_title))
            elements.append(Paragraph(vuln.description, body_text))
            
            # Technical details section
            elements.append(Paragraph("<b>Technical Details</b>", section_title))
            tech_details = """The vulnerability was identified during the assessment through a combination of manual testing 
            and automated scanning. Exploitation of this vulnerability requires specific conditions and access levels 
            as detailed in the description."""
            elements.append(Paragraph(tech_details, body_text))
            
            # Impact section with styled formatting
            elements.append(Paragraph("<b>Business Impact</b>", section_title))
            if vuln.impact:
                elements.append(Paragraph(vuln.impact, body_text))
            else:
                elements.append(Paragraph("Impact analysis was not specified for this vulnerability.", body_text))
            
            # Remediation section with actionable steps
            elements.append(Paragraph("<b>Recommended Remediation</b>", section_title))
            if vuln.remediation:
                elements.append(Paragraph(vuln.remediation, body_text))
            else:
                elements.append(Paragraph("Specific remediation steps were not provided for this vulnerability.", body_text))
            
            # Risk assessment scoring
            elements.append(Paragraph("<b>Risk Assessment</b>", section_title))
            risk_factors = [
                ["Risk Factor", "Value", "Comments"],
                ["Likelihood", severity_to_value(vuln.severity), "Based on ease of exploitation"],
                ["Impact", severity_to_value(vuln.severity), "Based on potential damage"],
                ["Overall Risk", vuln.severity, "Combined assessment"]
            ]
            
            risk_table = Table(risk_factors, colWidths=[120, 80, 250])
            risk_table.setStyle(TableStyle([
                # Header row styling to match template exactly
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0B5345')),  # Dark green header matching template exactly
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),  # White text in header
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),  # Exact font size from template
                ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                ('TOPPADDING', (0, 0), (-1, 0), 8),
                # Data rows styling to match template
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#E8F5E9')),  # Light green background matching template
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#138D75')),  # Green grid color to match template exactly
                ('FONTSIZE', (0, 1), (-1, -1), 9),  # Exact font size from template
                ('BOTTOMPADDING', (0, 1), (-1, -1), 6),
                ('TOPPADDING', (0, 1), (-1, -1), 6),
            ]))
            elements.append(risk_table)
            
            # Add vulnerability images if available with better presentation
            vuln_images = [img for img in images if img.vulnerability_id == vuln.id]
            if vuln_images:
                elements.append(Paragraph("<b>Evidence and Screenshots</b>", section_title))
                elements.append(Paragraph("The following screenshots demonstrate the vulnerability:", body_text))
                
            for img in vuln_images:
                try:
                    img_reader = io.BytesIO(img.data)
                    image = RLImage(img_reader, width=450, height=320)
                    elements.append(image)
                    if img.description:
                        elements.append(Paragraph(f"<i>Figure: {img.description}</i>", styles['ImageCaption']))
                except Exception as e:
                    elements.append(Paragraph(f"Error loading image: {str(e)}", styles['Normal']))
            
            elements.append(Spacer(1, 20))
        
        elements.append(Spacer(1, 20))
    
    elements.append(PageBreak())
    
    # Conclusion
    elements.append(Paragraph("5. Conclusion", styles['Heading1']))
    conclusion_text = f"""
    Based on the assessment of {client.name}'s {report.domain_type.lower()} systems, a total of {stats['total']} 
    vulnerabilities were identified. These findings highlight areas that require attention to improve the overall 
    security posture.
    
    It is recommended that {client.name} address the identified vulnerabilities according to their severity, 
    beginning with critical and high-risk issues. Regular security assessments should be conducted to ensure 
    that security controls remain effective over time.
    """
    elements.append(Paragraph(conclusion_text, styles['Normal']))
    
    # Tools Used
    elements.append(Paragraph("6. Tools Used", styles['Heading1']))
    tools_text = """
    The following tools were used during the assessment:
    
    - Nmap: Network discovery and security auditing
    - Burp Suite: Web application security testing
    - OWASP ZAP: Web application vulnerability scanner
    - Metasploit: Penetration testing framework
    - Nessus: Vulnerability scanner
    - SQLmap: SQL injection detection and exploitation
    - Wireshark: Network protocol analyzer
    """
    elements.append(Paragraph(tools_text, styles['Normal']))
    
    # Build the PDF document
    doc.build(elements)
    
    # Get the PDF data and return it
    pdf_data = buffer.getvalue()
    buffer.close()
    return pdf_data

def get_severity_color(severity):
    """Return the appropriate color for each severity level"""
    severity_colors = {
        'Critical': 'danger',
        'High': 'danger',
        'Medium': 'warning',
        'Low': 'info'
    }
    return severity_colors.get(severity, 'secondary')

def severity_to_value(severity):
    """Convert severity level to a numeric value for risk assessment"""
    severity_values = {
        'Critical': 'High (9-10)',
        'High': 'Medium-High (7-8)',
        'Medium': 'Medium (4-6)',
        'Low': 'Low (1-3)'
    }
    return severity_values.get(severity, 'Not Rated')

def image_to_base64(image_data):
    """Convert image binary data to base64 string for HTML embedding"""
    try:
        if image_data:
            encoded = base64.b64encode(image_data).decode('utf-8')
            return f"data:image/jpeg;base64,{encoded}"
    except Exception:
        pass
    return None

def generate_docx_report(report, vulnerabilities, client, stats, images):
    """Generate an editable Word document report that exactly matches the template"""
    doc = Document()
    
    # Set margins to match template (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title Page - Exactly matching the provided template
    title = doc.add_heading("Vulnerability Assessment", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading("and Penetration Testing", 0)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add spacing before the subtitles
    for _ in range(4):
        doc.add_paragraph("")
    
    app_report = doc.add_heading(f"{report.domain_type} Application Security Report", 0)
    app_report.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    l1_report = doc.add_heading("L1 Report", 1)
    l1_report.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date = doc.add_heading(f"{datetime.now().strftime('%d %B %Y')}", 1)
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Client Details on a separate page immediately after the title page
    doc.add_heading("Client Details", 1)
    client_table = doc.add_table(rows=3, cols=2)
    client_table.style = 'Table Grid'
    
    # Populate the client details table
    client_table.rows[0].cells[0].text = "Organization"
    client_table.rows[0].cells[1].text = client.name
    
    client_table.rows[1].cells[0].text = "Contact Person"
    client_table.rows[1].cells[1].text = client.contact_email.split("@")[0] if client.contact_email else ""
    
    client_table.rows[2].cells[0].text = "Email id"
    client_table.rows[2].cells[1].text = client.contact_email if client.contact_email else ""
    
    # Document Revision
    doc.add_heading("Document Revision", 2)
    revision_table = doc.add_table(rows=3, cols=2)
    revision_table.style = 'Table Grid'
    
    # Populate the revision table
    revision_table.rows[0].cells[0].text = "Version"
    revision_table.rows[0].cells[1].text = f"{report.version}.0"
    
    revision_table.rows[1].cells[0].text = "Date"
    revision_table.rows[1].cells[1].text = datetime.now().strftime('%d %B %Y')
    
    revision_table.rows[2].cells[0].text = "Submitted By"
    revision_table.rows[2].cells[1].text = client.submitted_by or "Security Team"
    
    # Test Performed Details
    doc.add_heading("Test Performed Details", 2)
    test_table = doc.add_table(rows=4, cols=2)
    test_table.style = 'Table Grid'
    
    # Populate the test details table
    test_table.rows[0].cells[0].text = "Testing done By"
    test_table.rows[0].cells[1].text = client.testing_done_by or "Security Team"
    
    test_table.rows[1].cells[0].text = "Reviewed By"
    test_table.rows[1].cells[1].text = client.reviewed_by or "Senior Security Analyst"
    
    test_table.rows[2].cells[0].text = "Date"
    test_table.rows[2].cells[1].text = f"{report.start_date.strftime('%d %B %Y')} - {report.end_date.strftime('%d %B %Y')}"
    
    test_table.rows[3].cells[0].text = "Version"
    test_table.rows[3].cells[1].text = f"{report.version}.0"
    
    # Add page break after client information section
    doc.add_page_break()
    
    # Engagement Overview - Match template exactly
    doc.add_heading("Engagement Overview", 1)
    engagement_text = f"""{client.name} has engaged with CyHEX Infotech to conduct a penetration test of their {report.domain_type}
Application. This report contains all the results of the report as well as all the action items that
were included in the penetration test. The purpose of this report is to present the current
security level of the external perimeters including gaps, vulnerabilities, and misconfigurations.
The findings presented in this report should be fixed to improve the security level of the
network systems."""
    doc.add_paragraph(engagement_text)
    
    # Service Description
    doc.add_heading("Service Description", 2)
    service_text = """Web application Vulnerability Assessment and Penetration Testing (VAPT) is the process of
simulating real-world attacks by using the same techniques as malicious hackers. For a
security assessment that goes beyond a simple vulnerability scanner, you need experts in the
industry. Scrut Automation conducts its penetration test by approaching the scope with both
a manual and automatic approach."""
    doc.add_paragraph(service_text)
    
    # Web Application Penetration Test section
    doc.add_heading("Web Application Penetration Test", 2)
    webapp_text = """Our application-level penetration testing consists of both unauthenticated and authenticated
testing using both automated and manual methods with particular emphasis placed on
identifying vulnerabilities associated with the OWASP Top 10 Most Critical Application
Vulnerabilities. It is important to note that a penetration test is not just an automated
vulnerability scan, and a large portion of web application penetration testing is a manual
process with a skilled engineer attempting to identify, exploit, and evaluate the associate risk
of security issues."""
    doc.add_paragraph(webapp_text)
    
    # Project Objectives
    doc.add_heading("Project Objectives", 2)
    objectives_text = """CyHEX Infotech conduct all testing manually combined with custom and commercial tools that
perform unique attack approaches on the network to make sure we cover the whole system in
the test. Our expert knowledge and experience are the value we provide in our services"""
    doc.add_paragraph(objectives_text)
    
    # Add a page break before Table of Contents
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading("Table of Contents", 1)
    toc_items = [
        "1. Executive Summary",
        "    1.1 Summary",
        "    1.2 Approach",
        "    1.3 Disclaimer",
        "    1.4 Limitations",
        "    1.5 OWASP TOP 10",
        "    1.6 Vulnerability Scoring",
        "2. Checklist",
        "3. Scope",
        "    3.1 Key Findings",
        "    3.2 Vulnerability Graph",
        "4. Findings",
        "5. Conclusions",
        "6. Tools Used"
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading("1. Executive Summary", 1)
    doc.add_heading("1.1 Summary", 2)
    
    summary_text = f"""Our security team conducted a penetration test on {client.name}'s {report.domain_type}
environment, starting on {report.start_date.strftime('%d %B %Y')}. This assessment, combining automated tools and
manual checks, aimed to uncover technical weaknesses in the application that
could be exploited. The report details the identified vulnerabilities, their severity,
and provides recommendations to mitigate any security risks they might pose.

A total of {stats['total']} vulnerabilities were identified, categorized as follows:
- Critical: {stats['critical']}
- High: {stats['high']}
- Medium: {stats['medium']}
- Low: {stats['low']}"""
    doc.add_paragraph(summary_text)
    
    # Approach
    doc.add_heading("1.2 Approach", 2)
    
    # Add introductory text
    doc.add_paragraph("The approach for this assessment included the following activities:")
    
    # Add each bullet point as a separate paragraph for proper spacing
    bullet_points = [
        "Exploring various application functionalities to enumerate threat & vulnerability in alignment with Open Web Application Security Project (OWASP) Top 10 vulnerabilities.",
        "Performing information gathering/fingerprinting to identify software used/its version, web server details, ports, and services open, etc.",
        "Performing vulnerability scanning to identify common vulnerabilities in the application layer and by using Burp and various testing tools in the Kali Linux distribution in conjunction with a range of manual analysis.",
        "It should be noted that customized payloads and attack vectors were configured in Burp Suite to further enhance the identification of weakness in the application.",
        "Analyzing the automated scan results for any vulnerabilities and ease of exploitability and providing proof of concept where safe exploits are possible.",
        "Post-Exploitation process will be performed once we get access to the device using identified vulnerabilities/exploits.",
        "Reporting identified vulnerabilities and recommended solutions to mitigate them; for ease of mitigation activities for application support personnel/ developers' further details of CWEs were added."
    ]
    
    # Add each bullet point with proper spacing
    for point in bullet_points:
        paragraph = doc.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(point)
        # Add reduced spacing for bullet points while still ensuring each appears on a new line
        paragraph.paragraph_format.space_after = Pt(10)
        paragraph.paragraph_format.line_spacing = 1.2
    
    # Disclaimer
    doc.add_heading("1.3 Disclaimer", 2)
    disclaimer_text = """This report contains confidential information about the security of the client's systems. 
Distribution of this report should be limited to authorized personnel only.
The testing was performed during a limited timeframe and may not have identified all possible vulnerabilities."""
    doc.add_paragraph(disclaimer_text)
    
    # Limitations
    doc.add_heading("1.4 Limitations", 2)
    doc.add_paragraph(report.limitations or "Not specified")
    
    # OWASP TOP 10
    doc.add_heading("1.5 OWASP TOP 10", 2)
    
    # Add introductory text
    doc.add_paragraph("The OWASP Top 10 is a regularly updated report outlining security concerns for web application security. The top 10 vulnerabilities for 2021 are:")
    
    # OWASP Categories
    owasp_categories = [
        "A01:2021 – Broken Access Control",
        "A02:2021 – Cryptographic Failures",
        "A03:2021 – Injection",
        "A04:2021 – Insecure Design",
        "A05:2021 – Security Misconfiguration",
        "A06:2021 – Vulnerable and Outdated Components",
        "A07:2021 – Identification and Authentication Failures",
        "A08:2021 – Software and Data Integrity Failures",
        "A09:2021 – Security Logging and Monitoring Failures",
        "A10:2021 – Server-Side Request Forgery"
    ]
    
    # Add each bullet point with proper spacing
    for category in owasp_categories:
        paragraph = doc.add_paragraph()
        paragraph.style = 'List Bullet'
        paragraph.add_run(category)
        # Add reduced spacing for bullet points while still ensuring each appears on a new line
        paragraph.paragraph_format.space_after = Pt(10)
        paragraph.paragraph_format.line_spacing = 1.2
    
    # Vulnerability Scoring
    doc.add_heading("1.6 Vulnerability Scoring", 2)
    doc.add_paragraph("The Risk level is divided in four categories:")
    
    # Create a 2-column table for vulnerability severity descriptions that matches the template exactly
    severity_table = doc.add_table(rows=5, cols=2)
    severity_table.style = 'Table Grid'
    
    # Set the header row
    header_cells = severity_table.rows[0].cells
    header_cells[0].text = "Severity"
    header_cells[1].text = "DESCRIPTION"
    
    # Set the color of the header row - dark green header exactly matching template
    for cell in header_cells:
        cell_shading = parse_xml(f'''
        <w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
        w:fill="0B5345" w:val="clear"/>
        ''')
        cell._element.tcPr.append(cell_shading)
        # Set text color to white
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True
    
    # Add data rows exactly matching the template text with proper wrapping
    severity_rows = [
        ("Critical", "Critical vulnerabilities provide attackers with remote root or administrator capabilities. Malicious users have the ability to compromise the entire host. Easy to detect and exploit and result in large asset damage."),
        ("High", "Exploitation of the vulnerability discovered on the system can directly lead an attacker to information allowing them to gain privileged access (e.g., administrator or root) to the system. These issues are often difficult to detect and exploit but can result in large asset damage."),
        ("Medium", "The vulnerability discovered on the system can directly lead to an attacker gaining non-privileged access (e.g., as a standard user) to the system or the vulnerability provides access that can be leveraged within one step to gain administrator-level access. These issues are easy to detect and exploit, but typically result in small asset damage."),
        ("Low", "The vulnerability discovered on the system provides low-level, but sufficient data to the attacker that may be used to launch a more informed attack against the target environment. In addition, the vulnerability may indirectly lead to an attacker gaining some form of access to the system. These issues can be difficult to detect and exploit and typically result in small asset damage.")
    ]
    
    # Set the width of each cell to match the template
    for cell in header_cells:
        # Set column widths (approximate values)
        if cell.text == "Severity":
            cell.width = Inches(1.2)  # Narrower for severity column
        else:
            cell.width = Inches(4.0)  # Wider for description column
    
    for i, (severity, description) in enumerate(severity_rows, start=1):
        row = severity_table.rows[i].cells
        row[0].text = severity
        row[1].text = description
        
        # Apply the same column widths to the data rows
        row[0].width = Inches(1.2)
        row[1].width = Inches(4.0)
        
        # Set the color of severity cells exactly matching the template
        severity_colors = {
            "Critical": "FF0000",  # Red
            "High": "FFA500",      # Orange
            "Medium": "FFFF00",    # Yellow
            "Low": "7FFF00"        # Green
        }
        
        # Apply text color to ensure visibility
        text_colors = {
            "Critical": "FFFFFF",  # White text for dark background
            "High": "000000",      # Black text
            "Medium": "000000",    # Black text
            "Low": "000000"        # Black text
        }
        
        if severity in severity_colors:
            # Apply severity cell background color
            cell_shading = parse_xml(f'''
            <w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
            w:fill="{severity_colors[severity]}" w:val="clear"/>
            ''')
            row[0]._element.tcPr.append(cell_shading)
            
            # Apply text color to ensure visibility
            for paragraph in row[0].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor.from_string(text_colors[severity])
                    run.bold = True
            
            # Set cell background to light green for descriptions
            cell_shading = parse_xml(f'''
            <w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
            w:fill="E8F5E9" w:val="clear"/>
            ''')
            row[1]._element.tcPr.append(cell_shading)
    
    doc.add_page_break()
    
    # Checklist
    doc.add_heading("2. Checklist", 1)
    doc.add_paragraph("The following security testing checklist was used during the assessment:")
    
    # Add the security testing checklist table
    if report.checklist_data:
        try:
            checklist_data = json.loads(report.checklist_data)
            
            # Create checklist table
            checklist_table = doc.add_table(rows=1, cols=2)
            checklist_table.style = 'Table Grid'
            
            # Set column widths
            for cell in checklist_table.columns[0].cells:
                cell.width = Inches(4.0)
            for cell in checklist_table.columns[1].cells:
                cell.width = Inches(2.0)
            
            # Add header row
            header_cells = checklist_table.rows[0].cells
            header_cells[0].text = "Category"
            header_cells[1].text = "Outcome"
            
            # Style the header row
            for cell in checklist_table.rows[0].cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                shading_elm = parse_xml(f'<w:shd {{http://schemas.openxmlformats.org/wordprocessingml/2006/main}} w:fill="0B5345"/>')
                cell._tc.get_or_add_tcPr().append(shading_elm)
                # Add white text color
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            
            # Add data rows
            for category in checklist_data["categories"]:
                row_cells = checklist_table.add_row().cells
                row_cells[0].text = category["name"]
                row_cells[1].text = category["outcome"]
                # Center align the outcome cell
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add light green background to data cells
                for cell in row_cells:
                    shading_elm = parse_xml(f'<w:shd {{http://schemas.openxmlformats.org/wordprocessingml/2006/main}} w:fill="E8F5E9"/>')
                    cell._tc.get_or_add_tcPr().append(shading_elm)
        except Exception as e:
            logging.error(f"Error processing checklist data for DOCX: {str(e)}")
            doc.add_paragraph("Error loading checklist data.")
    else:
        doc.add_paragraph("No checklist data available.")
    
    doc.add_page_break()
    
    # Scope
    doc.add_heading("3. Scope", 1)
    doc.add_paragraph(report.scope or "Not specified")
    
    # Key Findings
    doc.add_heading("3.1 Key Findings", 2)
    
    # Create findings table
    findings_table = doc.add_table(rows=6, cols=2)
    findings_table.style = 'Table Grid'
    
    # Headers
    header_cells = findings_table.rows[0].cells
    header_cells[0].text = "Severity"
    header_cells[1].text = "Count"
    
    # Populate data
    findings_data = [
        ["Critical", str(stats['critical'])],
        ["High", str(stats['high'])],
        ["Medium", str(stats['medium'])],
        ["Low", str(stats['low'])],
        ["Total", str(stats['total'])]
    ]
    
    for i, (severity, count) in enumerate(findings_data, 1):
        row = findings_table.rows[i].cells
        row[0].text = severity
        row[1].text = count
    
    # Vulnerability Graph section
    doc.add_heading("3.2 Vulnerability Graph", 2)
    doc.add_paragraph("This section contains a graphical representation of the vulnerabilities in the report.")
    
    doc.add_page_break()
    
    # Findings
    doc.add_heading("4. Findings", 1)
    
    # Organize vulnerabilities into groups
    grouped_vulns = {}
    ungrouped_vulns = []
    
    for vuln in vulnerabilities:
        if vuln.group_id:
            group_id = vuln.group_id
            if group_id not in grouped_vulns:
                grouped_vulns[group_id] = {
                    "name": vuln.group.name,
                    "vulns": []
                }
            grouped_vulns[group_id]["vulns"].append(vuln)
        else:
            ungrouped_vulns.append(vuln)
    
    section_index = 1
    
    # Display grouped vulnerabilities
    for group_id, group_data in grouped_vulns.items():
        doc.add_heading(f"4.{section_index} {group_data['name']}", 2)
        doc.add_paragraph(f"This group contains {len(group_data['vulns'])} related vulnerabilities")
        
        # Sort vulnerabilities within group by severity
        severity_order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}
        sorted_vulns = sorted(group_data['vulns'], key=lambda x: severity_order.get(x.severity, 4))
        
        for i, vuln in enumerate(sorted_vulns, 1):
            doc.add_heading(f"4.{section_index}.{i} {vuln.title} ({vuln.severity})", 3)
            
            if vuln.cwe_id:
                doc.add_paragraph(f"CWE ID: {vuln.cwe_id}")
            
            doc.add_heading("Description:", 4)
            doc.add_paragraph(vuln.description)
            
            doc.add_heading("Impact:", 4)
            doc.add_paragraph(vuln.impact or "Not specified")
            
            doc.add_heading("Remediation:", 4)
            doc.add_paragraph(vuln.remediation or "Not specified")
            
            # Add vulnerability images if available
            vuln_images = [img for img in images if img.vulnerability_id == vuln.id]
            if vuln_images:
                doc.add_heading("Evidence:", 4)
                
                for img in vuln_images:
                    try:
                        # Save image to a temporary file
                        temp_image = tempfile.NamedTemporaryFile(delete=False, suffix='.jpg')
                        temp_image.write(img.data)
                        temp_image.close()
                        
                        # Add the image to the document
                        doc.add_picture(temp_image.name, width=Inches(4))
                        
                        # Add description if available
                        if img.description:
                            p = doc.add_paragraph(img.description)
                            p.italic = True
                        
                        # Clean up temp file
                        os.unlink(temp_image.name)
                    except Exception as e:
                        doc.add_paragraph(f"Error loading image: {str(e)}")
        
        section_index += 1
    
    # Display ungrouped vulnerabilities
    severity_order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3}
    sorted_ungrouped = sorted(ungrouped_vulns, key=lambda x: severity_order.get(x.severity, 4))
    
    for vuln in sorted_ungrouped:
        doc.add_heading(f"4.{section_index} {vuln.title} ({vuln.severity})", 2)
        
        if vuln.cwe_id:
            doc.add_paragraph(f"CWE ID: {vuln.cwe_id}")
        
        doc.add_heading("Description:", 3)
        doc.add_paragraph(vuln.description)
        
        doc.add_heading("Impact:", 3)
        doc.add_paragraph(vuln.impact or "Not specified")
        
        doc.add_heading("Remediation:", 3)
        doc.add_paragraph(vuln.remediation or "Not specified")
        
        # Add vulnerability images if available
        vuln_images = [img for img in images if img.vulnerability_id == vuln.id]
        if vuln_images:
            doc.add_heading("Evidence:", 3)
            
            for img in vuln_images:
                try:
                    # Save image to a temporary file
                    temp_image = tempfile.NamedTemporaryFile(delete=False, suffix='.jpg')
                    temp_image.write(img.data)
                    temp_image.close()
                    
                    # Add the image to the document
                    doc.add_picture(temp_image.name, width=Inches(4))
                    
                    # Add description if available
                    if img.description:
                        p = doc.add_paragraph(img.description)
                        p.italic = True
                    
                    # Clean up temp file
                    os.unlink(temp_image.name)
                except Exception as e:
                    doc.add_paragraph(f"Error loading image: {str(e)}")
        
        section_index += 1
    
    doc.add_page_break()
    
    # Conclusion
    doc.add_heading("5. Conclusion", 1)
    conclusion_text = f"""Based on the assessment of {client.name}'s {report.domain_type.lower()} systems, a total of {stats['total']} 
vulnerabilities were identified. These findings highlight areas that require attention to improve the overall 
security posture.

It is recommended that {client.name} address the identified vulnerabilities according to their severity, 
beginning with critical and high-risk issues. Regular security assessments should be conducted to ensure 
that security controls remain effective over time."""
    doc.add_paragraph(conclusion_text)
    
    # Tools Used
    doc.add_heading("6. Tools Used", 1)
    tools_text = """The following tools were used during the assessment:

- Nmap: Network discovery and security auditing
- Burp Suite: Web application security testing
- OWASP ZAP: Web application vulnerability scanner
- Metasploit: Penetration testing framework
- Nessus: Vulnerability scanner
- SQLmap: SQL injection detection and exploitation
- Wireshark: Network protocol analyzer"""
    doc.add_paragraph(tools_text)
    
    # Save to memory stream
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_data = docx_buffer.getvalue()
    docx_buffer.close()
    
    return docx_data
